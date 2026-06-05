#!/usr/bin/env python3
"""
题库文件转换工具
支持 docx / xlsx / csv / txt / pdf -> JSON（刷题宝格式）

用法:
    python convert_to_json.py <文件路径> [--name 题库名称] [--output 输出路径]
    python convert_to_json.py 题库.docx
    python convert_to_json.py 题库.xlsx --name "期末复习"
    python convert_to_json.py 题库.pdf -o output.json
"""

import argparse
import json
import re
import sys
import csv
import os
from pathlib import Path
from typing import Optional


# ============================================================
# 文本解析核心（docx / txt / pdf 共用）
# ============================================================

def parse_text_to_questions(text: str) -> list[dict]:
    """从纯文本中解析题目列表。支持多种常见格式。"""

    lines = text.split('\n')
    lines = [l.strip() for l in lines if l.strip()]

    # 按题目分组：遇到新题号就切分
    blocks = []
    current_block = []

    for line in lines:
        # 跳过标题行（如 "一、 单项选择题"）
        if re.match(r'^[一二三四五六七八九十]+[、.．]', line):
            continue
        # 跳过纯数字序号行（如 "一、"）
        if re.match(r'^[一二三四五六七八九十]+、\s*$', line):
            continue

        # 检测新题号开始: 1. / 1、/ 1) / 第1题 等（不要求后面有空格）
        is_question_start = bool(re.match(r'^(?:第\s*)?\d+[.、．\)）]', line))

        if is_question_start and current_block:
            blocks.append(current_block)
            current_block = [line]
        elif is_question_start:
            current_block = [line]
        else:
            current_block.append(line)

    if current_block:
        blocks.append(current_block)

    questions = []
    for block in blocks:
        q = parse_block_v2(block)
        if q and len(q['question']) >= 4 and q['answer'] != ['']:
            questions.append(q)

    return questions


def parse_block_v2(lines: list[str]) -> Optional[dict]:
    """
    解析一个题目的文本块（多行）。
    格式示例:
      1、题目内容（ ）。\t正确答案：A
      A、 选项A
      B、 选项B
      C、 选项C
      D、 选项D
    """
    if not lines:
        return None

    # 第一行是题干（可能包含答案，用 tab 分隔）
    first_line = lines[0]

    # 拆分 tab（题干和答案可能在同一行）
    parts = first_line.split('\t')
    question_part = parts[0].strip()
    answer_part = parts[1].strip() if len(parts) > 1 else ''

    # 去掉题号
    question_text = re.sub(r'^(?:第\s*)?\d+[.、．\)）]\s*', '', question_part).strip()
    if not question_text:
        return None

    options = {}
    answer_line = ''
    explanation_line = ''
    other_lines = []

    # 解析答案部分（tab 后面的内容）
    if answer_part:
        ans_match = re.search(
            r'(?:\(?正确\)?答案|答案|answer|答)\s*[：:]\s*(.+)',
            answer_part, re.I
        )
        if ans_match:
            answer_line = ans_match.group(1).strip()

    # 解析后续行（选项、答案、解析）
    for i in range(1, len(lines)):
        line = lines[i]

        # 选项行: A、/ A. / A：/ A:
        opt_match = re.match(r'^([A-Za-z])[.、．:：]\s*(.*)', line)
        if opt_match:
            options[opt_match.group(1).upper()] = opt_match.group(2).strip()
            continue

        # 答案行（如果 tab 里没有答案，可能单独一行）
        ans_match = re.search(
            r'(?:\(?正确\)?答案|答案|answer|答)\s*[：:]\s*(.+)',
            line, re.I
        )
        if ans_match and not answer_line:
            answer_line = ans_match.group(1).strip()
            continue

        # 解析行
        exp_match = re.match(r'^(解析|解释|说明|分析)\s*[：:]\s*(.*)', line)
        if exp_match:
            explanation_line = exp_match.group(2).strip()
            continue

        # 可能是题干的续行（被换行截断的情况）
        if not options and not answer_line:
            question_text += line
        else:
            other_lines.append(line)

    # 解析答案
    answer = parse_answer(answer_line)

    # 也尝试从题干中提取答案（某些格式答案嵌在题干里）
    if not answer:
        # 检查题干中是否有 "答案是X" 的模式
        inline_ans = re.search(r'答案[是为：:]\s*([A-Za-z]+)', question_text)
        if inline_ans:
            answer = parse_answer(inline_ans.group(1))

    if not question_text:
        return None

    # 清理题干
    question_text = clean_question(question_text)

    result = {
        'question': question_text,
        'answer': answer if answer else [''],
    }

    if options:
        result['options'] = options
    if explanation_line:
        result['explanation'] = explanation_line
    elif other_lines:
        result['explanation'] = '\n'.join(other_lines)

    return result


def merge_lines(lines: list[str]) -> list[str]:
    """
    智能合并被换行截断的行。
    规则：如果一行不以题号、选项字母、答案标记开头，则合并到上一行。
    """
    result = []
    for line in lines:
        stripped = line.strip()
        if not stripped:
            result.append('')
            continue

        # 检查是否是新题目的开始
        is_new_question = bool(re.match(r'^\d+[.、．\)）]', stripped))
        # 检查是否是选项行
        is_option = bool(re.match(r'^[A-Za-z][.、．:：\s]', stripped))
        # 检查是否是答案行
        is_answer = bool(re.match(r'^(\(?正确\)?答案|答案|answer|答)', stripped, re.I))
        # 检查是否是解析行
        is_explanation = bool(re.match(r'^(解析|解释|说明|分析)[：:]', stripped))

        if is_new_question or is_option or is_answer or is_explanation:
            result.append(stripped)
        elif result and result[-1]:
            # 合并到上一行
            result[-1] += stripped
        else:
            result.append(stripped)

    return result


def parse_single_block(block: str) -> Optional[dict]:
    """解析单个题目文本块。"""
    lines = [l.strip() for l in block.split('\n') if l.strip()]
    if not lines:
        return None

    # 提取题干（第一行，去掉题号）
    first_line = lines[0]
    question_text = re.sub(r'^(?:第\s*)?\d+[.、．\)）]\s*', '', first_line).strip()
    if not question_text:
        return None

    options = {}
    answer_line = ''
    explanation_line = ''
    other_lines = []

    for i in range(1, len(lines)):
        line = lines[i]
        # 选项: A. / A、/ A：/ A: / (A)
        opt_match = re.match(r'^([A-Za-z])[.、．:：（(]\s*(.*)', line)
        if opt_match:
            options[opt_match.group(1).upper()] = opt_match.group(2).rstrip('）)')
            continue

        # 答案行: 正确答案：X / 答案：X / Answer: X
        ans_match = re.match(
            r'^(?:\(?正确\)?答案|答案|answer|答)\s*[：:]\s*(.*)',
            line, re.I
        )
        if ans_match:
            answer_line = ans_match.group(1).strip()
            continue

        # 解析行
        exp_match = re.match(r'^(解析|解释|说明|分析)\s*[：:]\s*(.*)', line)
        if exp_match:
            explanation_line = exp_match.group(2).strip()
            continue

        other_lines.append(line)

    # 解析答案
    answer = parse_answer(answer_line)

    # 判断题型
    q_type = detect_question_type(question_text, options, answer)

    if not question_text:
        return None

    result = {
        'question': clean_question(question_text),
        'answer': answer if answer else [''],
    }

    if options:
        result['options'] = options
    if explanation_line:
        result['explanation'] = explanation_line
    elif other_lines:
        result['explanation'] = '\n'.join(other_lines)

    return result


def parse_answer(raw: str) -> list[str]:
    """解析答案字符串为列表。"""
    if not raw:
        return []

    cleaned = raw.strip()

    # 判断题
    if re.match(r'^(对|正确|√|T|true|是|错|不正确|错误|×|F|false|否)$', cleaned, re.I):
        if re.match(r'^(对|正确|√|T|true|是)$', cleaned, re.I):
            return ['true']
        else:
            return ['false']

    # 多选题: A、B、C 或 A,B,C 或 ABC 或 A B C
    # 先检查是否是连续字母（如 "ABCD"）
    if re.match(r'^[A-Z]{2,}$', cleaned):
        return sorted(list(cleaned))

    # 去掉常见分隔符
    parts = re.split(r'[、，,\s]+', cleaned)
    letters = []
    for p in parts:
        p = p.strip().upper()
        if re.match(r'^[A-Z]$', p):
            letters.append(p)

    if letters:
        return sorted(set(letters))

    # 如果没有匹配到字母，返回原文（可能是填空题）
    return [cleaned]


def detect_question_type(question: str, options: dict, answer: list[str]) -> str:
    """根据内容推断题型。"""
    # 判断题：答案是 true/false
    if len(answer) == 1 and answer[0] in ('true', 'false'):
        return 'judge'

    # 有选项的题
    if options:
        if len(answer) > 1:
            return 'multiple'
        return 'single'

    # 没有选项
    if '____' in question or '（  ）' in question or '（）' in question:
        return 'blank'

    return 'short'


def clean_question(text: str) -> str:
    """清理题干文本。"""
    # 去掉末尾的（ ）或 ( ) 占位符（可能有空格）
    text = re.sub(r'[（(]\s*[）)]\s*[。.]?\s*$', '', text)
    # 去掉中间的空括号占位符
    text = re.sub(r'[（(]\s*[）)]', '', text)
    # 去掉末尾残留的半括号
    text = re.sub(r'[（(]\s*$', '', text)
    # 去掉末尾多余标点
    text = text.rstrip('。.，,')
    # 处理 HTML 实体
    text = text.replace('&quot;', '"').replace('&amp;', '&').replace('&lt;', '<').replace('&gt;', '>')
    return text.strip()


# ============================================================
# DOCX 解析
# ============================================================

def extract_from_docx(filepath: str) -> tuple[str, list[str]]:
    """从 docx 文件提取文本。返回 (标题, 行列表)。"""
    try:
        from docx import Document
    except ImportError:
        print("错误: 需要安装 python-docx: pip install python-docx", file=sys.stderr)
        sys.exit(1)

    doc = Document(filepath)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    # 也检查表格中的内容
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    t = p.text.strip()
                    if t:
                        paragraphs.append(t)

    # 第一行可能是标题
    title = ''
    lines = paragraphs
    if lines and not re.match(r'^\d+[.、．\)）]', lines[0]):
        title = lines[0]
        lines = lines[1:]

    return title, lines


# ============================================================
# XLSX 解析
# ============================================================

def extract_from_xlsx(filepath: str) -> tuple[str, list[dict]]:
    """从 Excel 文件提取题目。"""
    try:
        import openpyxl
    except ImportError:
        print("错误: 需要安装 openpyxl: pip install openpyxl", file=sys.stderr)
        sys.exit(1)

    wb = openpyxl.load_workbook(filepath, read_only=True)
    ws = wb.active

    rows = list(ws.iter_rows(values_only=True))
    if not rows:
        return '', []

    # 检测表头
    header = [str(c).strip().lower() if c else '' for c in rows[0]]

    # 查找列索引
    col_map = detect_columns(header)

    title = Path(filepath).stem
    questions = []

    for row in rows[1:]:
        if not row or not any(row):
            continue
        q = parse_xlsx_row(row, col_map)
        if q:
            questions.append(q)

    return title, questions


def detect_columns(header: list[str]) -> dict:
    """智能检测 Excel 列映射。"""
    col_map = {'question': -1, 'answer': -1, 'explanation': -1}
    option_cols = {}  # A -> col_idx

    for i, h in enumerate(header):
        h = h.strip()
        if not h:
            continue

        # 题干列
        if h in ('题目', '题干', '问题', 'question', 'content', '题'):
            col_map['question'] = i
        # 答案列
        elif h in ('答案', '正确答案', 'answer', 'ans'):
            col_map['answer'] = i
        # 解析列
        elif h in ('解析', '解释', '说明', 'explanation'):
            col_map['explanation'] = i
        # 选项列: A / 选项A / option_a
        elif re.match(r'^[A-Za-z]$', h):
            option_cols[h.upper()] = i
        elif re.match(r'^选项\s*[A-Za-z]$', h):
            option_cols[h[-1].upper()] = i
        elif h.startswith('option') and len(h) > 6:
            letter = h[-1].upper()
            if letter.isalpha():
                option_cols[letter] = i

    col_map['options'] = option_cols
    return col_map


def parse_xlsx_row(row: tuple, col_map: dict) -> Optional[dict]:
    """解析 Excel 单行。"""
    row = list(row)
    q_text = ''
    answer_raw = ''
    explanation = ''
    options = {}

    # 提取题干
    if col_map['question'] >= 0 and col_map['question'] < len(row):
        q_text = str(row[col_map['question']] or '').strip()
    else:
        # 默认第一列是题干
        q_text = str(row[0] or '').strip() if row else ''

    if not q_text:
        return None

    # 清理题干（去掉题号）
    q_text = re.sub(r'^(?:第\s*)?\d+[.、．\)）]\s*', '', q_text).strip()

    # 提取选项
    for letter, idx in col_map.get('options', {}).items():
        if idx < len(row) and row[idx]:
            options[letter] = str(row[idx]).strip()

    # 提取答案
    if col_map['answer'] >= 0 and col_map['answer'] < len(row):
        answer_raw = str(row[col_map['answer']] or '').strip()

    # 提取解析
    if col_map['explanation'] >= 0 and col_map['explanation'] < len(row):
        explanation = str(row[col_map['explanation']] or '').strip()

    answer = parse_answer(answer_raw)

    result = {
        'question': clean_question(q_text),
        'answer': answer if answer else [''],
    }
    if options:
        result['options'] = options
    if explanation:
        result['explanation'] = explanation

    return result


# ============================================================
# CSV 解析
# ============================================================

def extract_from_csv(filepath: str) -> tuple[str, list[dict]]:
    """从 CSV 文件提取题目。"""
    # 尝试不同编码
    for encoding in ['utf-8-sig', 'utf-8', 'gbk', 'gb18030', 'latin-1']:
        try:
            with open(filepath, 'r', encoding=encoding) as f:
                content = f.read()
            break
        except (UnicodeDecodeError, UnicodeError):
            continue
    else:
        print("错误: 无法识别 CSV 文件编码", file=sys.stderr)
        sys.exit(1)

    # 自动检测分隔符
    sniffer = csv.Sniffer()
    try:
        dialect = sniffer.sniff(content[:2048])
    except csv.Error:
        dialect = csv.excel

    reader = csv.reader(content.splitlines(), dialect)
    rows = [row for row in reader if any(cell.strip() for cell in row)]

    if not rows:
        return '', []

    # 第一行可能是表头
    header = [c.strip().lower() for c in rows[0]]
    first_data_row = rows[1] if len(rows) > 1 else []

    # 检测是否有表头
    has_header = any(h in ('题目', '题干', '问题', 'question', '答案', 'answer') for h in header)

    title = Path(filepath).stem

    if has_header:
        col_map = detect_columns(header)
        questions = []
        for row in rows[1:]:
            q = parse_xlsx_row(row, col_map)
            if q:
                questions.append(q)
    else:
        # 无表头，当作纯文本处理
        text = '\n'.join(','.join(row) for row in rows)
        questions = parse_text_to_questions(text)

    return title, questions


# ============================================================
# PDF 解析
# ============================================================

def extract_from_pdf(filepath: str) -> tuple[str, list[str]]:
    """从 PDF 文件提取文本。"""
    try:
        import pdfplumber
    except ImportError:
        print("错误: 需要安装 pdfplumber: pip install pdfplumber", file=sys.stderr)
        sys.exit(1)

    paragraphs = []
    with pdfplumber.open(filepath) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                paragraphs.extend(text.split('\n'))

    title = ''
    lines = [l.strip() for l in paragraphs if l.strip()]
    if lines and not re.match(r'^\d+[.、．\)）]', lines[0]):
        title = lines[0]
        lines = lines[1:]

    return title, lines


# ============================================================
# TXT / Markdown 解析
# ============================================================

def extract_from_txt(filepath: str) -> tuple[str, list[str]]:
    """从纯文本文件提取。"""
    for encoding in ['utf-8-sig', 'utf-8', 'gbk', 'gb18030', 'latin-1']:
        try:
            with open(filepath, 'r', encoding=encoding) as f:
                content = f.read()
            break
        except (UnicodeDecodeError, UnicodeError):
            continue
    else:
        print("错误: 无法识别文件编码", file=sys.stderr)
        sys.exit(1)

    lines = [l.strip() for l in content.split('\n') if l.strip()]
    title = ''
    if lines and not re.match(r'^\d+[.、．\)）]', lines[0]):
        title = lines[0]
        lines = lines[1:]

    return title, lines


# ============================================================
# 主转换入口
# ============================================================

def convert_file(filepath: str, name: Optional[str] = None, output: Optional[str] = None) -> dict:
    """转换单个文件为 JSON 格式。返回结果 dict。"""
    path = Path(filepath)
    ext = path.suffix.lower()

    print(f"正在解析: {path.name} ({ext})")

    title = name or path.stem
    questions = []

    if ext == '.docx':
        title_text, lines = extract_from_docx(filepath)
        if name:
            title = name
        elif title_text:
            title = title_text
        text = '\n'.join(lines)
        questions = parse_text_to_questions(text)

    elif ext in ('.xlsx', '.xls'):
        title_text, qs = extract_from_xlsx(filepath)
        if name:
            title = name
        elif title_text:
            title = title_text
        questions = qs

    elif ext == '.csv':
        title_text, qs = extract_from_csv(filepath)
        if name:
            title = name
        elif title_text:
            title = title_text
        questions = qs

    elif ext in ('.txt', '.md'):
        title_text, lines = extract_from_txt(filepath)
        if name:
            title = name
        elif title_text:
            title = title_text
        text = '\n'.join(lines)
        questions = parse_text_to_questions(text)

    elif ext == '.pdf':
        title_text, lines = extract_from_pdf(filepath)
        if name:
            title = name
        elif title_text:
            title = title_text
        text = '\n'.join(lines)
        questions = parse_text_to_questions(text)

    else:
        print(f"错误: 不支持的文件格式 {ext}", file=sys.stderr)
        print("支持: .docx / .xlsx / .xls / .csv / .txt / .md / .pdf", file=sys.stderr)
        sys.exit(1)

    # 补充 type 字段（App 需要）
    for q in questions:
        if 'type' not in q:
            q['type'] = detect_question_type(
                q['question'],
                q.get('options', {}),
                q.get('answer', [])
            )
        if 'tags' not in q:
            q['tags'] = []

    # 构造输出
    result = {
        'name': title,
        'questions': questions
    }

    # 统计
    type_counts = {}
    for q in questions:
        t = q.get('type', 'unknown')
        type_counts[t] = type_counts.get(t, 0) + 1

    print(f"\n解析完成!")
    print(f"题库名称: {title}")
    print(f"题目总数: {len(questions)}")
    print(f"题型分布:")
    type_labels = {
        'single': '单选题', 'multiple': '多选题', 'judge': '判断题',
        'blank': '填空题', 'short': '简答题', 'unknown': '未知'
    }
    for t, count in sorted(type_counts.items()):
        print(f"  {type_labels.get(t, t)}: {count}")

    # 输出文件
    if not output:
        output = str(path.parent / f"{path.stem}_converted.json")

    with open(output, 'w', encoding='utf-8') as f:
        json.dump(result, f, ensure_ascii=False, indent=2)

    print(f"\n已保存到: {output}")
    return result


# ============================================================
# CLI
# ============================================================

def main():
    parser = argparse.ArgumentParser(
        description='题库文件转换工具 - 将 Word/Excel/CSV/PDF/TXT 转为刷题宝 JSON 格式',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例:
  python convert_to_json.py 题库.docx
  python convert_to_json.py 题库.xlsx --name "期末复习"
  python convert_to_json.py 题库.pdf -o output.json
  python convert_to_json.py *.docx          # 批量转换
        """
    )
    parser.add_argument('files', nargs='+', help='要转换的文件路径（支持多个）')
    parser.add_argument('--name', '-n', help='题库名称（默认使用文件名）')
    parser.add_argument('--output', '-o', help='输出文件路径（默认: 原文件名_converted.json）')
    parser.add_argument('--merge', '-m', action='store_true', help='合并多个文件为一个题库')

    args = parser.parse_args()

    if len(args.files) == 1:
        convert_file(args.files[0], args.name, args.output)
    elif args.merge:
        # 合并模式
        all_questions = []
        title = args.name or '合并题库'
        for f in args.files:
            print(f"\n{'='*40}")
            result = convert_file(f, name='__temp__')
            all_questions.extend(result['questions'])

        merged = {'name': title, 'questions': all_questions}
        output = args.output or 'merged_quiz.json'
        with open(output, 'w', encoding='utf-8') as out:
            json.dump(merged, out, ensure_ascii=False, indent=2)

        print(f"\n{'='*40}")
        print(f"合并完成! 共 {len(all_questions)} 题 -> {output}")
    else:
        # 批量模式：每个文件单独转换
        for f in args.files:
            print(f"\n{'='*40}")
            convert_file(f)


if __name__ == '__main__':
    main()
